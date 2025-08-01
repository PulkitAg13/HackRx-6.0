import logging
from docx import Document
from io import BytesIO
from typing import Dict, List
import xml.etree.ElementTree as ET
from backend.app.core.config import settings

logger = logging.getLogger(__name__)

class DocxParser:
    def extract_text(self, file_stream: BytesIO) -> Dict[str, any]:
        """
        Extract text and structure from DOCX file
        Returns:
            {
                "text": str,
                "metadata": {
                    "paragraphs": List[Dict],
                    "tables": List[Dict],
                    "headers": List[str],
                    "footers": List[str]
                }
            }
        """
        try:
            doc = Document(file_stream)
            result = {
                "text": "",
                "metadata": {
                    "paragraphs": [],
                    "tables": [],
                    "headers": [],
                    "footers": []
                }
            }
            
            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append({
                        "text": para.text,
                        "style": para.style.name
                    })
            result["metadata"]["paragraphs"] = paragraphs
            
            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)
            result["metadata"]["tables"] = tables
            
            # Extract headers and footers
            for section in doc.sections:
                for header in section.header.paragraphs:
                    if header.text.strip():
                        result["metadata"]["headers"].append(header.text)
                for footer in section.footer.paragraphs:
                    if footer.text.strip():
                        result["metadata"]["footers"].append(footer.text)
            
            # Combine all text
            all_text = [p["text"] for p in paragraphs]
            all_text.extend([" | ".join(row) for table in tables for row in table])
            result["text"] = "\n".join(all_text)
            
            return result
        except Exception as e:
            logger.error(f"DOCX parsing failed: {str(e)}")
            raise

def parse_docx(file_stream: BytesIO) -> Dict[str, any]:
    return DocxParser().extract_text(file_stream)